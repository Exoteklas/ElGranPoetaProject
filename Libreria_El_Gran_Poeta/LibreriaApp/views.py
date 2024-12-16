from django.http import HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.decorators import user_passes_test
from django.contrib.auth import login as auth_login
from django.contrib.auth import authenticate, login as auth_login
from .forms import LoginForm
from LibreriaApp.forms import BodegaForm, LibroForm, AutorForm
from LibreriaApp.models import Bodega, Libro, Autor,AuditLog
from django.contrib import messages
from django.urls import reverse_lazy
from django.contrib.auth.views import LogoutView
from docx import Document
from datetime import date
from soporteApp.models import TicketSoporte


# Create your views here.

#Todo lo que es login

@login_required
def jefe_bodega_view(request):
    return render(request, 'jefe_bodega.html')

@login_required
def bodeguero_view(request):
    # Obtener todos los libros de la base de datos
    libros = Libro.objects.all()
    
    # Crear el contexto para pasar a la plantilla
    context = {
        'libros': libros
    }
    
    # Renderizar la plantilla 'bodeguero.html' con los libros
    return render(request, 'bodeguero.html', context)

def group_required(group_name):
    def in_group(user):
        return user.groups.filter(name=group_name).exists() or user.is_superuser
    return user_passes_test(in_group)

@login_required
@group_required('Jefe de bodega')
def jefe_bodegas_view(request):
    return render(request, 'jefe_bodega.html')

@login_required
@group_required('Bodeguero')
def bodeguero_view(request):
    return render(request, 'bodeguero.html')


@login_required
def soporte_tickets(request):
    # Filtrar los tickets por usuario si es necesario, o todos los tickets
    tickets = TicketSoporte.objects.all()  # o .filter(usuario=request.user) si solo quieres mostrar los del usuario actual
    return render(request, 'soporte_tickets.html', {'tickets': tickets})

def login_view(request):
    error_message = None  
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, username=username, password=password)
            if user is not None:
                auth_login(request, user)
                # Redirigir según el grupo
                if user.groups.filter(name='Jefe de bodega').exists():
                    return redirect('jefe_bodega')
                elif user.groups.filter(name='Bodeguero').exists():
                    return redirect('bodeguero')
                elif user.groups.filter(name='Soporte').exists():
                    return redirect('soporte_tickets')
                else:
                    return redirect('login')  
            else:
                error_message = "Usuario o contraseña incorrectos."
    else:
        form = LoginForm()
    
    return render(request, 'login.html', {'form': form, 'error_message': error_message})


class CustomLogoutView(LogoutView):
    next_page = reverse_lazy('login')


#Aca empiezan las cosas de la base

def jefecrud(request):
    return render(request, 'jefe_bodega.html')

def listadoBodega(request):
    bodegas = Bodega.objects.all()
    data = {'bodegas': bodegas}
    return render(request, 'CRUD/bodegas.html', data)

def agregarBodega(request):
    form = BodegaForm()
    if request.method == 'POST':
        form = BodegaForm(request.POST)
        if form.is_valid():
            nueva_bodega = form.save()
            # Aquí se registra el movimiento en el log de auditoría
            AuditLog.objects.create(
                action='Agregar Bodega',
                model_name='Bodega',  # El nombre del modelo que estás utilizando
                user=request.user.username,  # Guarda el nombre de usuario
                group=request.user.groups.first().name if request.user.groups.exists() else 'Sin grupo'  # Guarda el nombre del grupo
            )
            return redirect('listado_bodega')
    data = {'form': form}
    return render(request, 'CRUD/agregarBodega.html', data)

def eliminarBodega(request, id):
    bodega = Bodega.objects.get(id=id)
    libros = Libro.objects.filter(bodega=bodega)

    if libros.exists() and libros.filter(stock__gt=0).exists():
        messages.error(request, 'No puedes eliminar la bodega porque tiene libros con stock.')
    else:
        # Aquí se registra el movimiento en el log de auditoría
        AuditLog.objects.create(
            action='Eliminar Bodega',
            model_name='Bodega',
            user=request.user.username,
            group=request.user.groups.first().name if request.user.groups.exists() else 'Sin grupo'  # Asegúrate de que el usuario tenga un grupo
        )
        bodega.delete()
        messages.success(request, 'Bodega eliminada exitosamente.')
    
    return redirect('listado_bodega')

    

def actualizarBodega(request, id):
    bodega = Bodega.objects.get(id=id)
    form = BodegaForm(instance=bodega)
    if request.method == 'POST':
        form = BodegaForm(request.POST, instance=bodega)
        if form.is_valid():
            form.save()
            # Aquí se registra el movimiento en el log de auditoría
            AuditLog.objects.create(
                action='Actualizar Bodega',
                model_name='Bodega',
                user=request.user.username,
                group=request.user.groups.first().name if request.user.groups.exists() else 'Sin grupo'  # Asegúrate de que el usuario tenga un grupo
            )
            messages.success(request, 'Bodega actualizada exitosamente.')
            return redirect('listado_bodega')
    
    data = {'form': form}
    return render(request, 'CRUD/agregarBodega.html', data)



#Libros
def listadoLibros(request):
    libros = Libro.objects.all()
    data = {'libros': libros}
    return render(request, 'CRUD/libros.html', data)


def agregarLibros(request):
    form = LibroForm()
    if request.method == 'POST':
        form = LibroForm(request.POST)
        if form.is_valid():
            nuevo_libro = form.save()
            # Registro en el log de auditoría
            AuditLog.objects.create(
                action='Agregar Libro',
                model_name='Libro',
                user=request.user.username,
                group=request.user.groups.first().name if request.user.groups.exists() else 'Sin grupo'
            )
            return redirect('listado_libros')
    data = {'form': form}
    return render(request, 'CRUD/agregarLibros.html', data)

def eliminarLibro(request, id):
    libro = Libro.objects.get(id=id)
    libro.delete()
    # Registro en el log de auditoría
    AuditLog.objects.create(
        action='Eliminar Libro',
        model_name='Libro',
        user=request.user.username,
        group=request.user.groups.first().name if request.user.groups.exists() else 'Sin grupo'
    )
    return redirect('listado_libros')

def actualizarLibros(request, id):
    libro = Libro.objects.get(id=id)
    form = LibroForm(instance=libro)
    if request.method == 'POST':
        form = LibroForm(request.POST, instance=libro)
        if form.is_valid():
            form.save()
            # Registro en el log de auditoría
            AuditLog.objects.create(
                action='Actualizar Libro',
                model_name='Libro',
                user=request.user.username,
                group=request.user.groups.first().name if request.user.groups.exists() else 'Sin grupo'
            )
            return redirect('listado_libros')
    data = {'form': form}
    return render(request, 'CRUD/agregarLibros.html', data)




#Autor
def listadoAutores(request):
    autores = Autor.objects.all()
    data = {'autores': autores}
    return render(request, 'CRUD/autores.html', data)

def agregarAutores(request):
    form = AutorForm()
    if request.method == 'POST':
        form = AutorForm(request.POST)
        if form.is_valid():
            nuevo_autor = form.save()
            # Registro en el log de auditoría
            AuditLog.objects.create(
                action='Agregar Autor',
                model_name='Autor',
                user=request.user.username,
                group=request.user.groups.first().name if request.user.groups.exists() else 'Sin grupo'
            )
            return redirect('listado_autores')
    data = {'form': form}
    return render(request, 'CRUD/agregarAutores.html', data)


def eliminarAutores(request, id):
    autor = Autor.objects.get(id=id)
    libros_asociados = Libro.objects.filter(autor=autor)

    if libros_asociados.exists() and libros_asociados.filter(stock__gt=0).exists():
        messages.error(request, 'No puedes eliminar este autor porque tiene libros asociados con stock.')
    else:
        autor.delete()
        messages.success(request, 'Autor eliminado exitosamente.')
        # Registra la eliminación exitosa en el log de auditoría
        AuditLog.objects.create(
            user=request.user,
            group=request.user.groups.first().name,
            action='Autor eliminado',
            model_name='Autor'
        )

    return redirect('listado_autores')


def actualizarAutores(request, id):
    autor = Autor.objects.get(id=id)
    form = AutorForm(instance=autor)
    if request.method == 'POST':
        form = AutorForm(request.POST, instance=autor)
        if form.is_valid():
            form.save()
            # Registro en el log de auditoría
            AuditLog.objects.create(
                action='Actualizar Autor',
                model_name='Autor',
                user=request.user.username,
                group=request.user.groups.first().name if request.user.groups.exists() else 'Sin grupo'
            )
            return redirect('listado_autores')
    data = {'form': form}
    return render(request, 'CRUD/agregarAutores.html', data)


def jefe_bodegas_view(request):
    bodegas = Bodega.objects.all()
    mensaje = "Por favor, elija una bodega para generar el informe:"
    context = {'bodegas': bodegas, 'mensaje': mensaje}
    return render(request, 'jefe_bodega.html', context)

def generar_informe(request):
    bodega_id = request.GET.get('bodega_id')
    if not bodega_id:
        return render(request, 'jefe_bodega.html', {'mensaje': 'Por favor, selecciona una bodega.'})

    documento = Document()

    bodega = get_object_or_404(Bodega, id=bodega_id)
    libros = Libro.objects.filter(bodega=bodega)

    documento.add_heading(f'Informe de Bodega {bodega.id}', 0)
    documento.add_paragraph(f'Ubicación de la bodega: {bodega.ubicacion}')
    documento.add_paragraph(f'Informe generado el {date.today()}')

    documento.add_paragraph('\n')

    documento.add_heading('Libros almacenados', level=1)
    documento.add_paragraph('\n')
    
    if libros.exists():
        tabla_libros = documento.add_table(rows=1, cols=3)
        tabla_libros.style = 'Table Grid'

        encabezados_libros = tabla_libros.rows[0].cells
        encabezados_libros[0].text = 'Nombre del libro'
        encabezados_libros[1].text = 'Autor'
        encabezados_libros[2].text = 'Stock'

        for libro in libros:
            fila = tabla_libros.add_row().cells
            fila[0].text = libro.nombre_libro
            fila[1].text = libro.autor.nombre
            fila[2].text = str(libro.stock)
    else:
        documento.add_paragraph('No hay libros almacenados en esta bodega.')
    documento.add_paragraph('\n')
   

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="informe_bodega_{bodega_id}.docx"'
    documento.save(response)

    return response

#BODEGUERO

# Lista de libros

def listadoLibrosBD(request):
    libros = Libro.objects.all()
    carrito_libros = []

    # Verifica si hay un carrito en la sesión
    if 'carrito' in request.session:
        # Obtén los datos de los libros y sus cantidades desde el carrito en sesión
        for libro_id, cantidad in request.session['carrito'].items():
            try:
                libro = Libro.objects.get(id=libro_id)
                carrito_libros.append({'libro': libro, 'cantidad': cantidad})
            except Libro.DoesNotExist:
                pass  # Si el libro no existe, ignóralo

    return render(request, 'CRUD/libros_bode.html', {'libros': libros, 'carrito_libros': carrito_libros})

# Agregar al carrito
def agregar_carrito(request, libro_id):
    libro = get_object_or_404(Libro, id=libro_id)

    # Verificar si el stock disponible es suficiente
    if libro.stock <= 0:
        return redirect('bodeguero')  # Redirigir si no hay stock

    # Inicializar el carrito en la sesión si no existe
    if 'carrito' not in request.session:
        request.session['carrito'] = {}

    carrito = request.session['carrito']

    # Verificar si el libro ya está en el carrito y actualizar la cantidad
    if str(libro.id) in carrito:
        if carrito[str(libro.id)] < libro.stock:  # Solo incrementar si hay stock disponible
            carrito[str(libro.id)] += 1
    else:
        # Si el libro no está en el carrito, agregarlo con cantidad 1
        carrito[str(libro.id)] = 1

    # Guardar el carrito actualizado en la sesión
    request.session['carrito'] = carrito

    return redirect('bodeguero')  # Redirigir de nuevo a la página del bodeguero

def mostrar_carrito(request):

    carrito = request.session.get('carrito', {})
    carrito_libros = []
    for libro_id, cantidad in carrito.items():
        libro = Libro.objects.get(id=libro_id)
        carrito_libros.append({
            'libro': libro,
            'cantidad': cantidad
        })

    return render(request, 'bodeguero.html', {'carrito_libros': carrito_libros})



def quitar_libro(request, libro_id):
    # Verificar si el carrito existe en la sesión
    if 'carrito' in request.session:
        carrito = request.session['carrito']

        # Eliminar el libro del carrito si está presente
        carrito.pop(str(libro_id), None)  # Elimina el libro por su id

        # Guardar el carrito actualizado en la sesión
        request.session['carrito'] = carrito

    return redirect('bodeguero')  # Redirigir a la página del bodeguero
    


def realizar_pedido(request):
    carrito = request.session.get('carrito', {})

    if not carrito:
        return redirect('bodeguero')  # Si el carrito está vacío, redirigir

    for libro_id, cantidad in carrito.items():
        libro = get_object_or_404(Libro, id=libro_id)

        if libro.stock >= cantidad:  # Asegúrate de que hay suficiente stock
            libro.stock -= cantidad
            libro.save()
        else:
            # Manejar el caso en que no hay suficiente stock
            # Aquí puedes mostrar un mensaje al usuario o redirigir
            # Para simplificar, solo eliminamos el libro del carrito
            del carrito[str(libro_id)]

    # Vaciar el carrito después de realizar el pedido
    request.session['carrito'] = {}

    return redirect('bodeguero')



